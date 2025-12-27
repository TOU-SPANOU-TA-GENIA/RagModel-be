# app/workflows/handlers/actions.py
"""
Action node handlers.
These nodes perform output actions like generating reports, sending emails, saving files.
"""

import os
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime

from app.workflows.engine import NodeResult, ExecutionContext
from app.workflows.storage import workflow_storage
from app.utils.logger import setup_logger

logger = setup_logger(__name__)


def _clean_llm_output(text: str) -> str:
    """
    Clean LLM output by removing chain-of-thought artifacts and prompt echoing.
    """
    if not text:
        return ""
    
    import re
    import json
    
    # Detect prompt echoing - if output starts with instruction-like text
    prompt_echo_patterns = [
        r"(?:^|\n)(?:Παρακαλώ,?\s*απαντ[ήέ]στε|Please,?\s*respond)[^.]*[.\n]",
        r'^Για την απάντηση.*?(?=\{|\Z)',
        r'^Στο JSON.*?(?=\{|\Z)',
        r'^Πρέπει να.*?(?=\{|\Z)',
        r'^ΑΠΑΝΤΗΣΗ.*?(?=\{|\Z)',
        r'^Η απάντηση.*?(?=\{|\Z)',
        r'^To answer.*?(?=\{|\Z)',
        r'^The response.*?(?=\{|\Z)',
    ]
    
    cleaned = text
    for pattern in prompt_echo_patterns:
        cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE | re.DOTALL)
    
    # Try to extract JSON - look for the LAST complete JSON object
    # This handles cases where the model outputs multiple JSON attempts
    json_objects = []
    brace_count = 0
    start_idx = -1
    
    for i, char in enumerate(cleaned):
        if char == '{':
            if brace_count == 0:
                start_idx = i
            brace_count += 1
        elif char == '}':
            brace_count -= 1
            if brace_count == 0 and start_idx >= 0:
                json_str = cleaned[start_idx:i+1]
                try:
                    parsed = json.loads(json_str)
                    json_objects.append(parsed)
                except json.JSONDecodeError:
                    pass
                start_idx = -1
    
    # Use the last valid JSON object (usually the final answer)
    if json_objects:
        parsed = json_objects[-1]
        result_parts = []
        
        # Build Greek summary from parsed data
        if parsed.get('summary'):
            result_parts.append(parsed['summary'])
        
        if parsed.get('anomalies_found') or parsed.get('found'):
            count = parsed.get('anomaly_count', parsed.get('count', 0))
            severity = parsed.get('severity', 'unknown')
            if not result_parts:
                result_parts.append(f"Εντοπίστηκαν {count} ανωμαλίες (Σοβαρότητα: {severity})")
        
        if parsed.get('findings') or parsed.get('issues'):
            findings = parsed.get('findings', parsed.get('issues', []))
            result_parts.append("\n\nΕυρήματα:")
            for finding in findings[:10]:
                if isinstance(finding, dict):
                    desc = finding.get('description', finding.get('type', ''))
                    risk = finding.get('risk', '')
                    result_parts.append(f"  • {desc}" + (f" ({risk})" if risk else ""))
                elif isinstance(finding, str):
                    result_parts.append(f"  • {finding}")
        
        if result_parts:
            return '\n'.join(result_parts)
    
    # Fallback: Remove obvious noise patterns
    patterns_to_remove = [
        r"<think>.*?</think>",
        r"(?:^|\n)\*?\*?(?:Παρακαλώ|Please),?\s*(?:απαντ[ήέ]στε|respond)[^.\n]*[.\n]?",
        r"\*\*Απάντηση:\*\*\s*",
        r"(?:^|\n)(?:Okay|Ok|Let me|First|Starting).*?(?=\n\n|\Z)",
        r"(?:^|\n)(?:Για|Στο|Πρέπει|Η απάντηση).*?(?=\n\n|\Z)",
        r"```[a-z]*\n?|```",
        r"^\.\.\.\s*$",
        r"^\}\s*$",
    ]
    
    for pattern in patterns_to_remove:
        cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE | re.MULTILINE | re.DOTALL)
    
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned).strip()
    
    return cleaned if len(cleaned) > 20 else text


async def handle_generate_report(
    config: Dict[str, Any],
    context: ExecutionContext
) -> NodeResult:
    """
    Generate a report document.
    
    Config:
        report_template: str - Template name or 'default'
        output_format: str - 'docx', 'pdf', 'html', 'markdown'
        output_path: str - Where to save the report
        filename_template: str - Filename pattern
        language: str - Report language
        include_sections: List[str] - Sections to include
    """
    output_format = config.get('output_format', 'docx')
    output_path = config.get('output_path', '')
    filename_template = config.get('filename_template', 'report_{timestamp}')
    language = config.get('language', 'greek')
    include_sections = config.get('include_sections', ['summary', 'details', 'anomalies', 'recommendations'])
    
    # Gather data from context
    summary = context.get('summary', '')
    analysis = context.get('analysis_results', {})
    anomalies = context.get('anomalies', [])
    source_files = context.get('source_files', [])
    critical_found = context.get('critical_found', False)
    critical_keywords = context.get('critical_keywords', [])
    
    # Anomaly detection results
    anomalies_found = context.get('anomalies_found', False)
    anomaly_count = context.get('anomaly_count', 0)
    severity = context.get('severity', 'low')
    findings = context.get('findings', [])
    
    # Clean the analysis text - remove LLM "thinking" artifacts
    raw_analysis = analysis.get('analysis', '')
    cleaned_analysis = _clean_llm_output(raw_analysis)
    
    # Get parsed data
    parsed = analysis.get('parsed', {})
    
    # Build report content
    report_sections = {}
    
    if 'summary' in include_sections:
        # Try to get parsed summary first
        if parsed.get('summary'):
            report_sections['summary'] = parsed['summary']
        elif summary:
            report_sections['summary'] = summary
        else:
            report_sections['summary'] = cleaned_analysis[:500]
    
    if 'details' in include_sections:
        # Filter out junk files
        JUNK_FILES = {'thumbs.db', 'desktop.ini', '.ds_store', '.gitkeep'}
        clean_files = [
            f.get('name') for f in source_files[:20] 
            if f.get('name') and f.get('name', '').lower() not in JUNK_FILES
        ]
        
        report_sections['details'] = {
            'file_count': len(clean_files),
            'files': clean_files,
            'full_analysis': cleaned_analysis,
            'critical_found': critical_found,
            'critical_keywords': critical_keywords,
            'anomalies_found': anomalies_found,
            'anomaly_count': anomaly_count,
            'severity': severity,
            'total_financial_risk': parsed.get('total_financial_risk', 'N/A')
        }
    
    # Include findings from LLM analysis as anomalies
    if 'anomalies' in include_sections:
        # Use findings from LLM analysis if available
        if findings:
            report_sections['anomalies'] = findings
        elif anomalies:
            report_sections['anomalies'] = anomalies
    
    if 'recommendations' in include_sections:
        recommendations = _generate_recommendations(findings or anomalies, language, critical_found, anomalies_found)
        # Add critical alert if detected
        if critical_found and critical_keywords:
            if language == 'greek':
                recommendations.insert(0, f"🚨 ΚΡΙΣΙΜΟ: Ανιχνεύθηκαν κρίσιμες λέξεις-κλειδιά: {', '.join(critical_keywords)}")
            else:
                recommendations.insert(0, f"🚨 CRITICAL: Detected critical keywords: {', '.join(critical_keywords)}")
        # Add anomaly alert if found
        if anomalies_found and anomaly_count > 0:
            if language == 'greek':
                recommendations.insert(0, f"⚠️ ΑΝΩΜΑΛΙΕΣ: Εντοπίστηκαν {anomaly_count} ανωμαλίες (Σοβαρότητα: {severity})")
            else:
                recommendations.insert(0, f"⚠️ ANOMALIES: Found {anomaly_count} anomalies (Severity: {severity})")
        report_sections['recommendations'] = recommendations
    
    # Generate filename with multiple template options
    now = datetime.utcnow()
    filename = filename_template
    filename = filename.replace('{timestamp}', now.strftime('%Y%m%d_%H%M%S'))
    filename = filename.replace('{date}', now.strftime('%Y-%m-%d'))
    filename = filename.replace('{datetime}', now.strftime('%Y%m%d_%H%M'))
    filename = filename.replace('{year}', now.strftime('%Y'))
    filename = filename.replace('{month}', now.strftime('%m'))
    filename = filename.replace('{day}', now.strftime('%d'))
    filename = f"{filename}.{output_format}"
    
    # Determine full output path
    if output_path:
        full_path = Path(output_path) / filename
    else:
        full_path = Path('data/outputs') / filename
    
    full_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Generate document
    try:
        if output_format == 'docx':
            await _generate_docx_report(full_path, report_sections, language)
        elif output_format == 'markdown':
            await _generate_markdown_report(full_path, report_sections, language)
        elif output_format == 'html':
            await _generate_html_report(full_path, report_sections, language)
        else:
            await _generate_markdown_report(
                full_path.with_suffix('.md'), 
                report_sections, 
                language
            )
        
        context.set('report_path', str(full_path))
        context.set('report_filename', filename)
        
        logger.info(f"📊 Report generated: {full_path}")
        
        return NodeResult(
            success=True,
            output={
                'report_path': str(full_path),
                'filename': filename,
                'format': output_format,
                'sections': list(report_sections.keys())
            }
        )
    
    except Exception as e:
        logger.error(f"Report generation failed: {e}")
        return NodeResult(
            success=False,
            error=str(e)
        )


def _generate_recommendations(anomalies: List[Dict], language: str, critical_found: bool = False, anomalies_found: bool = False) -> List[str]:
    """Generate recommendations based on detected anomalies/findings."""
    recommendations = []
    
    if not anomalies:
        # Only say "no anomalies" if critical_found is also false
        if not critical_found and not anomalies_found:
            if language == 'greek':
                recommendations.append("✅ Δεν εντοπίστηκαν ανωμαλίες. Συνεχίστε την κανονική παρακολούθηση.")
            else:
                recommendations.append("✅ No anomalies detected. Continue normal monitoring.")
        return recommendations
    
    # Check if these are LLM findings (have 'risk' field) or standard anomalies (have 'severity')
    if anomalies and isinstance(anomalies[0], dict):
        if 'risk' in anomalies[0]:
            # LLM findings format
            high_risk = [a for a in anomalies if a.get('risk') in ['high', 'critical']]
            medium_risk = [a for a in anomalies if a.get('risk') == 'medium']
            
            if language == 'greek':
                if high_risk:
                    recommendations.append(f"🔴 ΑΜΕΣΗ ΔΡΑΣΗ: {len(high_risk)} ευρήματα υψηλού κινδύνου:")
                    for finding in high_risk[:5]:
                        rec = finding.get('recommendation', finding.get('description', ''))
                        if rec:
                            recommendations.append(f"  • {rec[:100]}")
                
                if medium_risk:
                    recommendations.append(f"🟡 ΔΙΕΡΕΥΝΗΣΗ: {len(medium_risk)} ευρήματα μέτριου κινδύνου για επανεξέταση.")
            else:
                if high_risk:
                    recommendations.append(f"🔴 IMMEDIATE ACTION: {len(high_risk)} high-risk findings:")
                    for finding in high_risk[:5]:
                        rec = finding.get('recommendation', finding.get('description', ''))
                        if rec:
                            recommendations.append(f"  • {rec[:100]}")
                
                if medium_risk:
                    recommendations.append(f"🟡 INVESTIGATE: {len(medium_risk)} medium-risk findings for review.")
        else:
            # Standard anomalies format
            high_severity = [a for a in anomalies if a.get('severity') == 'high']
            medium_severity = [a for a in anomalies if a.get('severity') == 'medium']
            
            if language == 'greek':
                if high_severity:
                    recommendations.append(f"⚠️ ΑΜΕΣΗ ΔΡΑΣΗ: Εντοπίστηκαν {len(high_severity)} κρίσιμες ανωμαλίες που χρήζουν άμεσης διερεύνησης.")
                    for a in high_severity[:3]:
                        recommendations.append(f"  • Ελέγξτε: {a.get('description', a.get('field', 'Άγνωστο'))}")
                
                if medium_severity:
                    recommendations.append(f"📋 ΠΑΡΑΚΟΛΟΥΘΗΣΗ: {len(medium_severity)} ανωμαλίες μέτριας σοβαρότητας για επανεξέταση.")
            else:
                if high_severity:
                    recommendations.append(f"⚠️ IMMEDIATE ACTION: {len(high_severity)} critical anomalies require investigation.")
                if medium_severity:
                    recommendations.append(f"📋 MONITOR: {len(medium_severity)} medium severity anomalies for review.")
    
    return recommendations


async def _generate_docx_report(
    path: Path,
    sections: Dict,
    language: str
):
    """Generate a Word document report."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Αναφορά Ανάλυσης' if language == 'greek' else 'Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC'))
        
        # Critical alert banner if detected
        details = sections.get('details', {})
        if details.get('critical_found'):
            doc.add_paragraph()  # Spacing
            alert_para = doc.add_paragraph()
            alert_run = alert_para.add_run('🚨 ΚΡΙΣΙΜΗ ΑΝΑΦΟΡΑ - ΑΠΑΙΤΕΙΤΑΙ ΑΜΕΣΗ ΠΡΟΣΟΧΗ 🚨')
            alert_run.bold = True
            alert_run.font.size = Pt(14)
            alert_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if details.get('critical_keywords'):
                keywords_para = doc.add_paragraph()
                keywords_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                keywords_para.add_run(f"Κρίσιμες λέξεις: {', '.join(details['critical_keywords'])}")
            doc.add_paragraph()  # Spacing
        
        # Summary section
        if 'summary' in sections:
            doc.add_heading('Σύνοψη' if language == 'greek' else 'Summary', 1)
            summary_text = sections['summary']
            if isinstance(summary_text, str) and summary_text:
                doc.add_paragraph(summary_text)
        
        # Details section
        if 'details' in sections:
            doc.add_heading('Λεπτομέρειες' if language == 'greek' else 'Details', 1)
            doc.add_paragraph(f"Αρχεία που αναλύθηκαν: {details.get('file_count', 0)}")
            
            if details.get('files'):
                doc.add_paragraph("Αρχεία:")
                for f in details['files'][:10]:
                    if f:  # Skip None
                        doc.add_paragraph(f"  • {f}", style='List Bullet')
            
            if details.get('full_analysis'):
                doc.add_heading('Πλήρης Ανάλυση' if language == 'greek' else 'Full Analysis', 2)
                analysis_text = details['full_analysis']
                # Limit length for readability
                if len(analysis_text) > 3000:
                    analysis_text = analysis_text[:3000] + "\n\n[...περικοπή...]"
                doc.add_paragraph(analysis_text)
        
        # Anomalies/Findings section
        if 'anomalies' in sections and sections['anomalies']:
            doc.add_heading('Ευρήματα Ελέγχου' if language == 'greek' else 'Audit Findings', 1)
            
            anomaly_list = sections['anomalies']
            
            # Check if these are LLM findings (have 'type' field) or standard anomalies
            if anomaly_list and isinstance(anomaly_list[0], dict) and 'type' in anomaly_list[0]:
                # LLM findings format
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                headers = ['Τύπος', 'Περιγραφή', 'Κίνδυνος', 'Σύσταση'] if language == 'greek' else ['Type', 'Description', 'Risk', 'Recommendation']
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                
                for finding in anomaly_list[:20]:
                    if isinstance(finding, dict):
                        row = table.add_row().cells
                        row[0].text = str(finding.get('type', ''))[:30]
                        row[1].text = str(finding.get('description', ''))[:100]
                        row[2].text = str(finding.get('risk', finding.get('severity', '')))
                        row[3].text = str(finding.get('recommendation', ''))[:80]
                        
                        # Add evidence as sub-paragraph if present
                        if finding.get('evidence'):
                            evidence_para = doc.add_paragraph()
                            evidence_para.add_run(f"   Στοιχεία: {finding['evidence'][:150]}").italic = True
            else:
                # Standard anomalies format
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                headers = ['Πεδίο', 'Τιμή', 'Σοβαρότητα', 'Περιγραφή'] if language == 'greek' else ['Field', 'Value', 'Severity', 'Description']
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                
                for anomaly in anomaly_list[:20]:
                    if isinstance(anomaly, dict):
                        row = table.add_row().cells
                        row[0].text = str(anomaly.get('field', ''))
                        row[1].text = str(anomaly.get('value', ''))
                        row[2].text = str(anomaly.get('severity', ''))
                        row[3].text = str(anomaly.get('description', ''))[:50]
        
        # Recommendations section
        if 'recommendations' in sections:
            doc.add_heading('Συστάσεις' if language == 'greek' else 'Recommendations', 1)
            for rec in sections['recommendations']:
                doc.add_paragraph(rec, style='List Bullet')
        
        doc.save(path)
        
    except ImportError:
        # Fallback to markdown
        await _generate_markdown_report(path.with_suffix('.md'), sections, language)


async def _generate_markdown_report(
    path: Path,
    sections: Dict,
    language: str
):
    """Generate a Markdown report."""
    lines = []
    
    lines.append(f"# {'Αναφορά Ανάλυσης' if language == 'greek' else 'Analysis Report'}")
    lines.append(f"\n*{datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}*\n")
    
    if 'summary' in sections:
        lines.append(f"## {'Σύνοψη' if language == 'greek' else 'Summary'}\n")
        lines.append(sections['summary'])
        lines.append("")
    
    if 'anomalies' in sections and sections['anomalies']:
        lines.append(f"## {'Ανωμαλίες' if language == 'greek' else 'Anomalies'}\n")
        lines.append(f"| {'Πεδίο' if language == 'greek' else 'Field'} | {'Τιμή' if language == 'greek' else 'Value'} | {'Σοβαρότητα' if language == 'greek' else 'Severity'} |")
        lines.append("|-------|-------|----------|")
        
        for a in sections['anomalies'][:20]:
            lines.append(f"| {a.get('field', '')} | {a.get('value', '')} | {a.get('severity', '')} |")
        lines.append("")
    
    if 'recommendations' in sections:
        lines.append(f"## {'Συστάσεις' if language == 'greek' else 'Recommendations'}\n")
        for rec in sections['recommendations']:
            lines.append(f"- {rec}")
        lines.append("")
    
    path.write_text('\n'.join(lines), encoding='utf-8')


async def _generate_html_report(
    path: Path,
    sections: Dict,
    language: str
):
    """Generate an HTML report."""
    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{'Αναφορά Ανάλυσης' if language == 'greek' else 'Analysis Report'}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        h1 {{ color: #0066cc; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #0066cc; color: white; }}
        .high {{ color: red; font-weight: bold; }}
        .medium {{ color: orange; }}
    </style>
</head>
<body>
    <h1>{'Αναφορά Ανάλυσης' if language == 'greek' else 'Analysis Report'}</h1>
    <p><em>{datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}</em></p>
"""
    
    if 'summary' in sections:
        html += f"<h2>{'Σύνοψη' if language == 'greek' else 'Summary'}</h2>\n<p>{sections['summary']}</p>\n"
    
    if 'anomalies' in sections and sections['anomalies']:
        html += f"<h2>{'Ανωμαλίες' if language == 'greek' else 'Anomalies'}</h2>\n<table>\n"
        html += f"<tr><th>{'Πεδίο' if language == 'greek' else 'Field'}</th><th>{'Τιμή' if language == 'greek' else 'Value'}</th><th>{'Σοβαρότητα' if language == 'greek' else 'Severity'}</th></tr>\n"
        
        for a in sections['anomalies'][:20]:
            severity_class = a.get('severity', 'medium')
            html += f"<tr><td>{a.get('field', '')}</td><td>{a.get('value', '')}</td><td class='{severity_class}'>{a.get('severity', '')}</td></tr>\n"
        
        html += "</table>\n"
    
    if 'recommendations' in sections:
        html += f"<h2>{'Συστάσεις' if language == 'greek' else 'Recommendations'}</h2>\n<ul>\n"
        for rec in sections['recommendations']:
            html += f"<li>{rec}</li>\n"
        html += "</ul>\n"
    
    html += "</body></html>"
    
    path.write_text(html, encoding='utf-8')


async def handle_send_email(
    config: Dict[str, Any],
    context: ExecutionContext
) -> NodeResult:
    """
    Send an email notification.
    
    Config:
        recipients: List[str] - Email addresses
        subject_template: str - Subject with placeholders
        body_template: str - Body with placeholders
        attach_report: bool - Attach generated report
        priority: str - 'low', 'normal', 'high'
        skip_if_no_smtp: bool - Skip without error if SMTP not configured
    """
    recipients = config.get('recipients', [])
    subject_template = config.get('subject_template', 'Workflow Notification')
    body_template = config.get('body_template', 'Workflow completed.')
    attach_report = config.get('attach_report', False)
    skip_if_no_smtp = config.get('skip_if_no_smtp', True)  # Default to skip
    
    if not recipients:
        return NodeResult(
            success=False,
            error="No recipients specified"
        )
    
    # Get user SMTP settings
    smtp_config = config.get('smtp', {})
    
    # Check if SMTP is configured
    if not smtp_config.get('host') and skip_if_no_smtp:
        logger.warning("📧 SMTP not configured - skipping email send")
        subject = _render_template(subject_template, context)
        return NodeResult(
            success=True,
            output={
                'skipped': True,
                'reason': 'SMTP not configured',
                'recipients': recipients,
                'subject': subject
            }
        )
    
    # Build email content from context
    subject = _render_template(subject_template, context)
    body = _render_template(body_template, context)
    
    # Get report attachment if requested
    attachment_path = None
    if attach_report:
        attachment_path = context.get('report_path')
    
    try:
        await _send_email(
            recipients=recipients,
            subject=subject,
            body=body,
            attachment_path=attachment_path,
            smtp_config=smtp_config
        )
        
        logger.info(f"📧 Email sent to {len(recipients)} recipients")
        
        return NodeResult(
            success=True,
            output={
                'recipients': recipients,
                'subject': subject,
                'has_attachment': attachment_path is not None
            }
        )
    
    except Exception as e:
        if skip_if_no_smtp:
            logger.warning(f"📧 Email sending failed (non-critical): {e}")
            return NodeResult(
                success=True,
                output={
                    'skipped': True,
                    'reason': str(e),
                    'recipients': recipients,
                    'subject': subject
                }
            )
        else:
            logger.error(f"Email sending failed: {e}")
            return NodeResult(
                success=False,
                error=str(e)
            )


def _render_template(template: str, context: ExecutionContext) -> str:
    """Render a template string with context values."""
    result = template
    
    # Get analysis results
    analysis = context.get('analysis_results', {})
    parsed = analysis.get('parsed', {})
    
    # Replace common placeholders
    replacements = {
        '{workflow_id}': context.workflow_id,
        '{execution_id}': context.execution_id,
        '{timestamp}': datetime.utcnow().strftime('%Y-%m-%d %H:%M'),
        '{date}': datetime.utcnow().strftime('%Y-%m-%d'),
        '{anomaly_count}': str(context.get('anomaly_count', 0)),
        '{severity}': str(context.get('severity', 'unknown')),
        '{summary}': parsed.get('summary', context.get('summary', ''))[:300],
        '{file_count}': str(len(context.get('source_files', []))),
        '{critical_keywords}': ', '.join(context.get('critical_keywords', [])),
        '{total_financial_risk}': str(parsed.get('total_financial_risk', 'N/A')),
        '{report_path}': str(context.get('report_path', '')),
    }
    
    for placeholder, value in replacements.items():
        result = result.replace(placeholder, str(value))
    
    return result


async def _send_email(
    recipients: List[str],
    subject: str,
    body: str,
    attachment_path: Optional[str] = None,
    smtp_config: Dict = None
):
    """Send an email using SMTP."""
    smtp_config = smtp_config or {}
    
    smtp_host = smtp_config.get('host', os.getenv('SMTP_HOST', 'localhost'))
    smtp_port = smtp_config.get('port', int(os.getenv('SMTP_PORT', '587')))
    smtp_user = smtp_config.get('username', os.getenv('SMTP_USERNAME'))
    smtp_pass = smtp_config.get('password', os.getenv('SMTP_PASSWORD'))
    smtp_tls = smtp_config.get('use_tls', True)
    from_addr = smtp_config.get('from_address', smtp_user or 'noreply@example.com')
    
    # Create message
    msg = MIMEMultipart()
    msg['From'] = from_addr
    msg['To'] = ', '.join(recipients)
    msg['Subject'] = subject
    
    msg.attach(MIMEText(body, 'plain', 'utf-8'))
    
    # Add attachment
    if attachment_path and Path(attachment_path).exists():
        with open(attachment_path, 'rb') as f:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(f.read())
        
        encoders.encode_base64(part)
        part.add_header(
            'Content-Disposition',
            f'attachment; filename="{Path(attachment_path).name}"'
        )
        msg.attach(part)
    
    # Send
    with smtplib.SMTP(smtp_host, smtp_port) as server:
        if smtp_tls:
            server.starttls()
        if smtp_user and smtp_pass:
            server.login(smtp_user, smtp_pass)
        server.send_message(msg)


async def handle_save_to_folder(
    config: Dict[str, Any],
    context: ExecutionContext
) -> NodeResult:
    """
    Save files to a destination folder.
    
    Config:
        destination_path: str - Where to save
        filename_template: str - Filename pattern
        overwrite: bool - Overwrite existing
        create_subfolders: bool - Create date-based subfolders
    """
    destination_path = config.get('destination_path', '')
    filename_template = config.get('filename_template', '{original_name}_{timestamp}')
    overwrite = config.get('overwrite', False)
    create_subfolders = config.get('create_subfolders', True)
    
    if not destination_path:
        return NodeResult(
            success=False,
            error="No destination path specified"
        )
    
    dest = Path(destination_path)
    
    # Create date subfolder if requested
    if create_subfolders:
        date_folder = datetime.utcnow().strftime('%Y-%m-%d')
        dest = dest / date_folder
    
    dest.mkdir(parents=True, exist_ok=True)
    
    saved_files = []
    errors = []
    
    # Save report if generated
    report_path = context.get('report_path')
    if report_path and Path(report_path).exists():
        try:
            import shutil
            report_name = Path(report_path).name
            target = dest / report_name
            
            if not target.exists() or overwrite:
                shutil.copy2(report_path, target)
                saved_files.append(str(target))
        except Exception as e:
            errors.append({'file': report_path, 'error': str(e)})
    
    context.set('saved_files', saved_files)
    
    logger.info(f"💾 Saved {len(saved_files)} files to {dest}")
    
    return NodeResult(
        success=len(errors) == 0,
        output={
            'saved_count': len(saved_files),
            'destination': str(dest),
            'files': saved_files,
            'errors': errors
        }
    )
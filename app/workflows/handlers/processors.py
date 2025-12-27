# app/workflows/handlers/processors.py
"""
Processor node handlers.
These nodes transform and analyze data.
"""

import json
import statistics
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime

from app.workflows.engine import NodeResult, ExecutionContext
from app.workflows.storage import workflow_storage
from app.utils.logger import setup_logger

logger = setup_logger(__name__)


async def handle_extract_content(
    config: Dict[str, Any],
    context: ExecutionContext
) -> NodeResult:
    """
    Extract content from files.
    
    Config:
        file_types: List[str] - Supported file types
        extract_tables: bool - Extract tables from documents
        extract_images: bool - Extract images (OCR)
        ocr_enabled: bool - Enable OCR for images
        language: str - OCR language
    
    Context Input:
        source_files: List[Dict] - Files to process
    
    Context Output:
        extracted_content: List[Dict] - Extracted content per file
    """
    source_files = context.get('source_files', [])
    if not source_files:
        return NodeResult(
            success=False,
            error="No source files provided"
        )
    
    # Filter out junk files
    JUNK_FILES = {'thumbs.db', 'desktop.ini', '.ds_store', '.gitkeep', '.gitignore'}
    source_files = [
        f for f in source_files 
        if f.get('name', '').lower() not in JUNK_FILES
    ]
    
    supported_types = config.get('file_types', ['pdf', 'docx', 'xlsx', 'txt', 'csv'])
    extract_tables = config.get('extract_tables', True)
    
    extracted = []
    errors = []
    
    for file_info in source_files:
        file_path = file_info.get('path')
        if not file_path:
            continue
        
        ext = Path(file_path).suffix.lower().lstrip('.')
        if ext not in supported_types:
            continue
        
        try:
            content = await _extract_file_content(file_path, ext, config)
            extracted.append({
                'file': file_info,
                'content': content,
                'extracted_at': datetime.utcnow().isoformat()
            })
        except Exception as e:
            errors.append({
                'file': file_path,
                'error': str(e)
            })
            logger.error(f"Failed to extract {file_path}: {e}")
    
    context.set('extracted_content', extracted)
    context.set('extraction_errors', errors)
    
    logger.info(f"📄 Extracted content from {len(extracted)} files")
    
    return NodeResult(
        success=len(extracted) > 0 or len(errors) == 0,
        output={
            'extracted_count': len(extracted),
            'error_count': len(errors),
            'errors': errors[:5]  # Limit errors in output
        }
    )


async def _extract_file_content(
    file_path: str,
    extension: str,
    config: Dict
) -> Dict[str, Any]:
    """Extract content from a single file."""
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    result = {
        'text': '',
        'tables': [],
        'metadata': {}
    }
    
    if extension == 'txt':
        result['text'] = path.read_text(encoding='utf-8', errors='ignore')
    
    elif extension == 'csv':
        import csv
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.DictReader(f)
            result['tables'] = [list(reader)]
            result['metadata']['columns'] = reader.fieldnames
    
    elif extension in ('xlsx', 'xls'):
        # Use openpyxl or similar
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True)
            tables = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                data = []
                for row in ws.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        data.append(list(row))
                if data:
                    tables.append({
                        'sheet': sheet,
                        'data': data
                    })
            result['tables'] = tables
            result['metadata']['sheets'] = wb.sheetnames
        except ImportError:
            result['text'] = f"[Excel file: {path.name}]"
    
    elif extension == 'pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(path)
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            result['text'] = '\n'.join(text_parts)
            result['metadata']['pages'] = len(doc)
            doc.close()
        except ImportError:
            result['text'] = f"[PDF file: {path.name}]"
    
    elif extension in ('docx', 'doc'):
        try:
            from docx import Document
            doc = Document(path)
            paragraphs = [p.text for p in doc.paragraphs]
            result['text'] = '\n'.join(paragraphs)
            
            # Extract tables
            if config.get('extract_tables', True):
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        table_data.append([cell.text for cell in row.cells])
                    result['tables'].append(table_data)
        except ImportError:
            result['text'] = f"[Word file: {path.name}]"
    
    return result


async def handle_llm_analysis(
    config: Dict[str, Any],
    context: ExecutionContext
) -> NodeResult:
    """
    Analyze content using LLM.
    
    Config:
        prompt_template: str - Prompt template with placeholders
        system_prompt: str - System prompt
        max_tokens: int - Max response tokens
        temperature: float - LLM temperature
        output_format: str - 'text', 'json', 'markdown'
        language: str - Response language
    
    Context Input:
        extracted_content: List[Dict] - Content to analyze (optional)
    
    Context Output:
        analysis_results: List[Dict] - Analysis per content
    """
    extracted_content = context.get('extracted_content', [])
    prompt_template = config.get('prompt_template', '')
    language = config.get('language', 'greek')
    
    # Build system prompt with language enforcement
    # default_system = 'Είσαι βοηθός ανάλυσης. ΑΠΑΝΤΗΣΕ ΜΟΝΟ ΣΤΑ ΕΛΛΗΝΙΚΑ. Μην γράφεις σκέψεις ή εξηγήσεις στα Αγγλικά.'
    from datetime import datetime
    current_date = datetime.now().strftime("%d/%m/%Y")
    default_system = f'Σημερινή ημερομηνία: {current_date}. Είσαι βοηθός ανάλυσης. ΑΠΑΝΤΗΣΕ ΜΟΝΟ ΣΤΑ ΕΛΛΗΝΙΚΑ.'
    if language != 'greek':
        default_system = 'You are an analysis assistant. Respond concisely without showing your thinking process.'
    system_prompt = config.get('system_prompt', default_system)
    
    # If no extracted content, use the prompt directly (for manual triggers)
    if not extracted_content:
        if not prompt_template:
            return NodeResult(
                success=False,
                error="No content to analyze and no prompt provided"
            )
        
        # Use prompt_template directly as the prompt
        prompt = prompt_template
        logger.info(f"🤖 Running direct LLM analysis with prompt: {prompt[:100]}...")
    else:
        # Build prompt from extracted content
        combined_text = []
        for item in extracted_content:
            content = item.get('content', {})
            file_info = item.get('file', {})
            
            text = content.get('text', '')
            tables = content.get('tables', [])
            
            # Format tables as text
            table_text = ''
            for table in tables:
                if isinstance(table, dict):
                    table_text += f"\n[Table from {table.get('sheet', 'unknown')}]\n"
                    for row in table.get('data', [])[:20]:
                        table_text += ' | '.join(str(cell) for cell in row) + '\n'
                elif isinstance(table, list):
                    for row in table[:20]:
                        table_text += ' | '.join(str(cell) for cell in row) + '\n'
            
            combined_text.append(f"=== {file_info.get('name', 'unknown')} ===\n{text}\n{table_text}")
        
        full_content = '\n\n'.join(combined_text)
        # Limit content aggressively for 6GB VRAM - 2000 chars default
        max_content = config.get('max_content_chars', 2000)
        original_len = len(full_content)
        if original_len > max_content:
            full_content = full_content[:max_content] + "\n[...]"
            logger.warning(f"Content truncated from {original_len} to {max_content} chars")
        
        prompt = prompt_template.replace('{content}', full_content) if prompt_template else f"Ανάλυσε:\n{full_content}"
    
    # Call LLM
    try:
        analysis = await _call_llm(prompt, system_prompt, config)
        
        # Critical keywords to detect (Greek military/security terms)
        CRITICAL_KEYWORDS = [
            'επίθεση', 'απειλή', 'κίνδυνος', 'όπλα', 'όπλο', 'εκρηκτικά', 
            'τρομοκρατ', 'βόμβα', 'πυρομαχικά', 'επείγον', 'κρίσιμο',
            'attack', 'threat', 'danger', 'weapon', 'explosive', 'bomb',
            'urgent', 'critical', 'emergency'
        ]
        
        # Anomaly keywords for supply chain detection
        ANOMALY_KEYWORDS = [
            'ανωμαλία', 'ανωμαλίες', 'υπερβολικ', 'ύποπτ', 'παράνομ',
            'απάτη', 'κλοπή', 'υπέρβαση', 'παρατυπία', 'παράβαση',
            'anomaly', 'suspicious', 'fraud', 'violation', 'overrun',
            'unauthorized', 'irregular', 'discrepancy'
        ]
        
        # Try to parse JSON response for structured data
        analysis_data = {'raw': analysis}
        critical_found = False
        found_keywords = []
        anomalies_found = False
        anomaly_count = 0
        severity = 'low'
        findings = []
        
        try:
            import json
            
            # Better JSON extraction - find all JSON objects and use the last valid one
            json_objects = []
            brace_count = 0
            start_idx = -1
            
            for i, char in enumerate(analysis):
                if char == '{':
                    if brace_count == 0:
                        start_idx = i
                    brace_count += 1
                elif char == '}':
                    brace_count -= 1
                    if brace_count == 0 and start_idx >= 0:
                        json_str = analysis[start_idx:i+1]
                        try:
                            parsed = json.loads(json_str)
                            # Only keep if it has relevant fields
                            if any(k in parsed for k in ['anomalies_found', 'found', 'anomaly_count', 'count', 'findings', 'issues', 'summary']):
                                json_objects.append(parsed)
                        except json.JSONDecodeError:
                            pass
                        start_idx = -1
            
            # Use the last valid JSON (usually the final answer)
            if json_objects:
                parsed = json_objects[-1]
                analysis_data.update(parsed)
                
                # Check for anomaly detection results (handle both field name variants)
                if parsed.get('anomalies_found') or parsed.get('found') or parsed.get('anomaly_count', 0) > 0 or parsed.get('count', 0) > 0:
                    anomalies_found = True
                    anomaly_count = parsed.get('anomaly_count', parsed.get('count', len(parsed.get('findings', parsed.get('issues', [])))))
                    severity = parsed.get('severity', 'medium')
                    findings = parsed.get('findings', parsed.get('issues', []))
                    
                    # High/critical severity anomalies trigger critical flag
                    if severity in ['high', 'critical'] or anomaly_count >= 3:
                        critical_found = True
                        found_keywords.append(f'{anomaly_count} ανωμαλίες')
                
                # Check parsed JSON for critical flags (intel scenario)
                if parsed.get('critical_found') or parsed.get('critical_keywords'):
                    critical_found = True
                    found_keywords.extend(parsed.get('critical_keywords', []))
                
                logger.info(f"✅ Parsed JSON: anomalies={anomaly_count}, severity={severity}")
                    
        except Exception as e:
            logger.warning(f"JSON parsing failed: {e}, using keyword fallback")
        
        # Fallback: Search for keywords in text
        if not critical_found and not anomalies_found:
            search_text = (prompt + ' ' + analysis).lower()
            
            # Check critical keywords
            for keyword in CRITICAL_KEYWORDS:
                if keyword.lower() in search_text:
                    critical_found = True
                    found_keywords.append(keyword)
            
            # Check anomaly keywords
            for keyword in ANOMALY_KEYWORDS:
                if keyword.lower() in search_text:
                    anomalies_found = True
                    if keyword not in found_keywords:
                        found_keywords.append(keyword)
            
            if found_keywords:
                logger.info(f"🚨 Keywords found via fallback: {found_keywords}")
        
        # Set context flags for both intel and anomaly scenarios
        context.set('critical_found', critical_found)
        context.set('critical_keywords', found_keywords)
        context.set('anomalies_found', anomalies_found)
        context.set('anomaly_count', anomaly_count)
        context.set('severity', severity)
        context.set('findings', findings)
        
        context.set('analysis_results', {
            'analysis': analysis,
            'parsed': analysis_data,
            'file_count': len(extracted_content),
            'critical_found': critical_found,
            'critical_keywords': found_keywords,
            'anomalies_found': anomalies_found,
            'anomaly_count': anomaly_count,
            'severity': severity,
            'findings': findings,
            'analyzed_at': datetime.utcnow().isoformat()
        })
        
        logger.info(f"🤖 LLM analysis completed - critical: {critical_found}, anomalies: {anomaly_count}")
        
        return NodeResult(
            success=True,
            output={
                'analysis_preview': analysis[:500] if analysis else '',
                'file_count': len(extracted_content),
                'critical_found': critical_found,
                'critical_keywords': found_keywords,
                'anomalies_found': anomalies_found,
                'anomaly_count': anomaly_count,
                'severity': severity
            }
        )
    
    except Exception as e:
        logger.error(f"LLM analysis failed: {e}")
        return NodeResult(
            success=False,
            error=str(e)
        )


async def _call_llm(prompt: str, system_prompt: str, config: Dict) -> str:
    """Call the LLM service for analysis."""
    try:
        from app.services.llm_service import llm_service
        
        # Limit max_tokens aggressively for 6GB VRAM GPU - 300 is enough for JSON
        max_tokens = min(config.get('max_tokens', 300), 400)
        
        # Default to Greek with no-thinking instruction
        if not system_prompt:
            system_prompt = 'Είσαι ελεγκτής. Απαντάς μόνο JSON.'
        
        logger.info(f"🤖 Calling LLM with {len(prompt)} char prompt, max_tokens={max_tokens}")
        
        response = await llm_service.generate(
            prompt=prompt,
            system_prompt=system_prompt,
            max_tokens=max_tokens,
            temperature=config.get('temperature', 0.1)  # Very low for consistent JSON
        )
        return response
    
    except ImportError as e:
        logger.warning(f"LLM service not available: {e}")
        return f"[LLM service not available - placeholder response for: {prompt[:100]}...]"
    
    except Exception as e:
        logger.error(f"LLM call failed: {e}")
        raise


async def handle_anomaly_detection(
    config: Dict[str, Any],
    context: ExecutionContext
) -> NodeResult:
    """
    Detect anomalies in data.
    
    Config:
        detection_type: str - 'statistical', 'pattern', 'llm'
        fields_to_monitor: List[str] - Fields to check
        threshold_std: float - Standard deviations for anomaly
        use_baseline: bool - Use stored baseline
        baseline_window_days: int - Days for baseline calculation
    
    Context Input:
        extracted_content: List[Dict] - Content with tables/data
    
    Context Output:
        anomalies: List[Dict] - Detected anomalies
    """
    extracted_content = context.get('extracted_content', [])
    detection_type = config.get('detection_type', 'statistical')
    threshold = config.get('threshold_std', 2.0)
    use_baseline = config.get('use_baseline', True)
    fields_to_monitor = config.get('fields_to_monitor', [])
    
    anomalies = []
    
    # Extract numeric data from tables
    numeric_data = _extract_numeric_data(extracted_content, fields_to_monitor)
    
    if detection_type == 'statistical':
        anomalies = await _statistical_anomaly_detection(
            numeric_data,
            threshold,
            use_baseline,
            context.workflow_id
        )
    elif detection_type == 'llm':
        anomalies = await _llm_anomaly_detection(
            extracted_content,
            config
        )
    
    context.set('anomalies', anomalies)
    context.set('anomaly_count', len(anomalies))
    
    # Update baseline with current data
    if use_baseline and numeric_data:
        await _update_baselines(numeric_data, context.workflow_id)
    
    logger.info(f"🔍 Anomaly detection found {len(anomalies)} anomalies")
    
    return NodeResult(
        success=True,
        output={
            'anomaly_count': len(anomalies),
            'anomalies': anomalies[:10],  # Limit output
            'detection_type': detection_type
        }
    )


def _extract_numeric_data(
    content: List[Dict],
    fields: List[str]
) -> Dict[str, List[float]]:
    """Extract numeric data from content."""
    data = {}
    
    for item in content:
        tables = item.get('content', {}).get('tables', [])
        
        for table in tables:
            # Handle different table formats
            if isinstance(table, dict):
                rows = table.get('data', [])
            else:
                rows = table
            
            if not rows:
                continue
            
            # First row is usually headers
            headers = rows[0] if rows else []
            
            for i, header in enumerate(headers):
                if header is None:
                    continue
                    
                header_str = str(header).lower().strip()
                
                # Check if this field should be monitored
                should_monitor = not fields or any(
                    f.lower() in header_str for f in fields
                )
                
                if should_monitor:
                    if header_str not in data:
                        data[header_str] = []
                    
                    # Extract numeric values from this column
                    for row in rows[1:]:
                        if i < len(row) and row[i] is not None:
                            try:
                                val = float(str(row[i]).replace(',', '.').replace(' ', ''))
                                data[header_str].append(val)
                            except (ValueError, TypeError):
                                pass
    
    return data


async def _statistical_anomaly_detection(
    data: Dict[str, List[float]],
    threshold: float,
    use_baseline: bool,
    workflow_id: str
) -> List[Dict]:
    """Statistical anomaly detection using z-scores."""
    anomalies = []
    
    for field, values in data.items():
        if len(values) < 3:
            continue
        
        # Get baseline or calculate from data
        baseline = None
        if use_baseline:
            baseline = workflow_storage.get_baseline(
                baseline_type='metric_avg',
                baseline_key=field,
                workflow_id=workflow_id
            )
        
        if baseline:
            mean = baseline['statistics'].get('mean', statistics.mean(values))
            std = baseline['statistics'].get('std', statistics.stdev(values))
        else:
            mean = statistics.mean(values)
            std = statistics.stdev(values) if len(values) > 1 else 0
        
        if std == 0:
            continue
        
        # Check each value
        for i, value in enumerate(values):
            z_score = abs(value - mean) / std
            if z_score > threshold:
                anomalies.append({
                    'field': field,
                    'value': value,
                    'expected_mean': round(mean, 2),
                    'z_score': round(z_score, 2),
                    'threshold': threshold,
                    'severity': 'high' if z_score > threshold * 1.5 else 'medium',
                    'description': f"{field}: τιμή {value} αποκλίνει σημαντικά από τον μέσο όρο {round(mean, 2)}"
                })
    
    return anomalies


async def _llm_anomaly_detection(
    content: List[Dict],
    config: Dict
) -> List[Dict]:
    """Use LLM for anomaly detection."""
    # Prepare content summary
    summary_parts = []
    for item in content:
        file_info = item.get('file', {})
        tables = item.get('content', {}).get('tables', [])
        
        for table in tables:
            if isinstance(table, dict):
                data = table.get('data', [])[:10]  # Sample
            else:
                data = table[:10]
            
            summary_parts.append(f"File: {file_info.get('name')}\nData sample: {data}")
    
    prompt = f"""Analyze the following data for anomalies, unusual patterns, or potential issues.
    
Data:
{chr(10).join(summary_parts[:5])}

Respond in JSON format with a list of anomalies:
[{{"field": "...", "description": "...", "severity": "high/medium/low"}}]
"""
    
    try:
        response = await _call_llm(prompt, "You are a data analyst. Detect anomalies.", config)
        
        # Parse JSON response
        import json
        try:
            anomalies = json.loads(response)
            if isinstance(anomalies, list):
                return anomalies
        except json.JSONDecodeError:
            pass
    except Exception as e:
        logger.error(f"LLM anomaly detection failed: {e}")
    
    return []


async def _update_baselines(
    data: Dict[str, List[float]],
    workflow_id: str
):
    """Update baselines with new data."""
    for field, values in data.items():
        if len(values) < 2:
            continue
        
        # Get existing baseline
        existing = workflow_storage.get_baseline(
            baseline_type='metric_avg',
            baseline_key=field,
            workflow_id=workflow_id
        )
        
        # Calculate new statistics
        new_stats = {
            'mean': statistics.mean(values),
            'std': statistics.stdev(values) if len(values) > 1 else 0,
            'min': min(values),
            'max': max(values),
            'count': len(values),
            'last_n_values': values[-10:]  # Keep last 10
        }
        
        # Merge with existing (exponential moving average)
        if existing:
            old_stats = existing['statistics']
            alpha = 0.3  # Weight for new data
            new_stats['mean'] = alpha * new_stats['mean'] + (1 - alpha) * old_stats.get('mean', new_stats['mean'])
            new_stats['std'] = alpha * new_stats['std'] + (1 - alpha) * old_stats.get('std', new_stats['std'])
            new_stats['count'] = old_stats.get('count', 0) + len(values)
        
        # Save
        workflow_storage.upsert_baseline(
            baseline_type='metric_avg',
            baseline_key=field,
            statistics=new_stats,
            workflow_id=workflow_id
        )


async def handle_cross_reference(
    config: Dict[str, Any],
    context: ExecutionContext
) -> NodeResult:
    """
    Cross-reference with existing knowledge.
    
    Config:
        reference_source: str - 'folder', 'database', 'index'
        reference_path: str - Path to reference data
        match_fields: List[str] - Fields to match on
        match_threshold: float - Fuzzy match threshold
    
    Context Output:
        matches: List[Dict] - Found matches
        contradictions: List[Dict] - Contradicting info
    """
    reference_source = config.get('reference_source', 'folder')
    reference_path = config.get('reference_path', '')
    
    # Get current analysis
    analysis = context.get('analysis_results', {})
    
    matches = []
    contradictions = []
    
    # Load reference data based on source
    if reference_source == 'folder':
        reference_data = await _load_reference_folder(reference_path)
    else:
        reference_data = []
    
    # TODO: Implement actual cross-referencing logic
    # This would use semantic similarity or keyword matching
    
    context.set('cross_reference_matches', matches)
    context.set('cross_reference_contradictions', contradictions)
    
    return NodeResult(
        success=True,
        output={
            'match_count': len(matches),
            'contradiction_count': len(contradictions)
        }
    )


async def _load_reference_folder(folder_path: str) -> List[Dict]:
    """Load reference documents from a folder."""
    # Placeholder - would load and index documents
    return []


async def handle_summarize(
    config: Dict[str, Any],
    context: ExecutionContext
) -> NodeResult:
    """
    Summarize content or analysis results.
    
    Config:
        summary_type: str - 'brief', 'detailed', 'bullets'
        max_length: int - Max summary length
        language: str - Output language
    """
    analysis = context.get('analysis_results', {})
    anomalies = context.get('anomalies', [])
    
    summary_type = config.get('summary_type', 'brief')
    language = config.get('language', 'greek')
    
    # Build summary prompt
    content_parts = []
    
    if analysis:
        content_parts.append(f"Analysis: {analysis.get('analysis', '')[:2000]}")
    
    if anomalies:
        anomaly_text = '\n'.join(
            f"- {a.get('description', a.get('field', 'Unknown'))}" 
            for a in anomalies[:10]
        )
        content_parts.append(f"Anomalies detected:\n{anomaly_text}")
    
    prompt = f"""Create a {summary_type} summary of the following:

{chr(10).join(content_parts)}

Summary language: {language}
"""
    
    try:
        summary = await _call_llm(
            prompt,
            f"You are a summarization assistant. Respond in {language}.",
            config
        )
        
        context.set('summary', summary)
        
        return NodeResult(
            success=True,
            output={
                'summary_preview': summary[:300] if summary else '',
                'summary_type': summary_type
            }
        )
    
    except Exception as e:
        return NodeResult(
            success=False,
            error=str(e)
        )
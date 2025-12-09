# app/services/file_generator.py
"""
File generator service for creating downloadable documents.
Supports DOCX, XLSX, PDF, TXT, MD, CSV.
"""

import io
import json
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple

from app.utils.logger import setup_logger

logger = setup_logger(__name__)


@dataclass
class GeneratedFile:
    """Represents a generated file."""
    filename: str
    mime_type: str
    data: bytes
    size_bytes: int
    
    def to_base64(self) -> str:
        import base64
        return base64.b64encode(self.data).decode('utf-8')


class FileGenerator:
    """Creates downloadable files from AI responses."""
    
    MIME_TYPES = {
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.pdf': 'application/pdf',
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.csv': 'text/csv',
        '.json': 'application/json',
        '.html': 'text/html',
    }
    
    @classmethod
    def create_docx(
        cls,
        content: str,
        title: Optional[str] = None,
        filename: Optional[str] = None
    ) -> GeneratedFile:
        """Create a Word document."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError("python-docx required: pip install python-docx")
        
        doc = Document()
        
        # Add title if provided
        if title:
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Split content by paragraphs
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                # Check for headers (# Markdown style)
                if para_text.startswith('# '):
                    doc.add_heading(para_text[2:], level=1)
                elif para_text.startswith('## '):
                    doc.add_heading(para_text[3:], level=2)
                elif para_text.startswith('### '):
                    doc.add_heading(para_text[4:], level=3)
                else:
                    doc.add_paragraph(para_text.strip())
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        data = buffer.getvalue()
        
        # Generate filename
        if not filename:
            safe_title = (title or "document").replace(' ', '_')[:30]
            filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return GeneratedFile(
            filename=filename,
            mime_type=cls.MIME_TYPES['.docx'],
            data=data,
            size_bytes=len(data)
        )
    
    @classmethod
    def create_xlsx(
        cls,
        data: List[List[Any]],
        headers: Optional[List[str]] = None,
        sheet_name: str = "Sheet1",
        filename: Optional[str] = None
    ) -> GeneratedFile:
        """Create an Excel spreadsheet."""
        try:
            import pandas as pd
        except ImportError:
            raise RuntimeError("pandas required: pip install pandas openpyxl")
        
        # Create DataFrame
        if headers:
            df = pd.DataFrame(data, columns=headers)
        else:
            df = pd.DataFrame(data)
        
        # Save to bytes
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        data_bytes = buffer.getvalue()
        
        if not filename:
            filename = f"spreadsheet_{datetime.now().strftime('%Y%m%d')}.xlsx"
        
        return GeneratedFile(
            filename=filename,
            mime_type=cls.MIME_TYPES['.xlsx'],
            data=data_bytes,
            size_bytes=len(data_bytes)
        )
    
    @classmethod
    def create_pdf(
        cls,
        content: str,
        title: Optional[str] = None,
        filename: Optional[str] = None
    ) -> GeneratedFile:
        """Create a PDF document."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
        except ImportError:
            raise RuntimeError("reportlab required: pip install reportlab")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        if title:
            story.append(Paragraph(title, styles['Title']))
            story.append(Spacer(1, 0.3 * inch))
        
        # Add content paragraphs
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                # Handle headers
                if para_text.startswith('# '):
                    story.append(Paragraph(para_text[2:], styles['Heading1']))
                elif para_text.startswith('## '):
                    story.append(Paragraph(para_text[3:], styles['Heading2']))
                else:
                    story.append(Paragraph(para_text.strip(), styles['Normal']))
                story.append(Spacer(1, 0.1 * inch))
        
        doc.build(story)
        data = buffer.getvalue()
        
        if not filename:
            safe_title = (title or "document").replace(' ', '_')[:30]
            filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d')}.pdf"
        
        return GeneratedFile(
            filename=filename,
            mime_type=cls.MIME_TYPES['.pdf'],
            data=data,
            size_bytes=len(data)
        )
    
    @classmethod
    def create_txt(
        cls,
        content: str,
        filename: Optional[str] = None
    ) -> GeneratedFile:
        """Create a plain text file."""
        data = content.encode('utf-8')
        
        if not filename:
            filename = f"document_{datetime.now().strftime('%Y%m%d')}.txt"
        
        return GeneratedFile(
            filename=filename,
            mime_type=cls.MIME_TYPES['.txt'],
            data=data,
            size_bytes=len(data)
        )
    
    @classmethod
    def create_md(
        cls,
        content: str,
        filename: Optional[str] = None
    ) -> GeneratedFile:
        """Create a Markdown file."""
        data = content.encode('utf-8')
        
        if not filename:
            filename = f"document_{datetime.now().strftime('%Y%m%d')}.md"
        
        return GeneratedFile(
            filename=filename,
            mime_type=cls.MIME_TYPES['.md'],
            data=data,
            size_bytes=len(data)
        )
    
    @classmethod
    def create_csv(
        cls,
        data: List[List[Any]],
        headers: Optional[List[str]] = None,
        filename: Optional[str] = None
    ) -> GeneratedFile:
        """Create a CSV file."""
        import csv
        
        buffer = io.StringIO()
        writer = csv.writer(buffer)
        
        if headers:
            writer.writerow(headers)
        writer.writerows(data)
        
        data_bytes = buffer.getvalue().encode('utf-8')
        
        if not filename:
            filename = f"data_{datetime.now().strftime('%Y%m%d')}.csv"
        
        return GeneratedFile(
            filename=filename,
            mime_type=cls.MIME_TYPES['.csv'],
            data=data_bytes,
            size_bytes=len(data_bytes)
        )
    
    @classmethod
    def create_from_type(
        cls,
        file_type: str,
        content: str,
        title: Optional[str] = None,
        data: Optional[List[List[Any]]] = None,
        headers: Optional[List[str]] = None,
        filename: Optional[str] = None
    ) -> GeneratedFile:
        """
        Create a file based on type string.
        
        Args:
            file_type: One of 'docx', 'xlsx', 'pdf', 'txt', 'md', 'csv'
            content: Text content for document types
            title: Optional title for documents
            data: List of rows for spreadsheet types
            headers: Optional headers for spreadsheet types
            filename: Optional custom filename
        """
        file_type = file_type.lower().strip('.')
        
        if file_type == 'docx':
            return cls.create_docx(content, title, filename)
        elif file_type == 'xlsx':
            if data is None:
                # Parse content as simple data
                data = [line.split('\t') for line in content.split('\n') if line.strip()]
            return cls.create_xlsx(data, headers, filename=filename)
        elif file_type == 'pdf':
            return cls.create_pdf(content, title, filename)
        elif file_type == 'txt':
            return cls.create_txt(content, filename)
        elif file_type == 'md':
            return cls.create_md(content, filename)
        elif file_type == 'csv':
            if data is None:
                data = [line.split(',') for line in content.split('\n') if line.strip()]
            return cls.create_csv(data, headers, filename)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")


# Convenience functions
def generate_document(
    content: str,
    file_type: str = 'docx',
    title: Optional[str] = None,
    filename: Optional[str] = None
) -> GeneratedFile:
    """Quick function to generate a document."""
    return FileGenerator.create_from_type(file_type, content, title, filename=filename)


def generate_spreadsheet(
    data: List[List[Any]],
    headers: Optional[List[str]] = None,
    file_type: str = 'xlsx',
    filename: Optional[str] = None
) -> GeneratedFile:
    """Quick function to generate a spreadsheet."""
    return FileGenerator.create_from_type(
        file_type, 
        "", 
        data=data, 
        headers=headers, 
        filename=filename
    )
# app/analysis/content_extractors.py
"""
Format-agnostic content extraction from various document types.
Supports PDF, DOCX, XLSX, images, and plain text.
"""

from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
import mimetypes
import hashlib

from app.utils.logger import setup_logger

logger = setup_logger(__name__)


@dataclass
class ExtractedContent:
    """Standardized container for extracted document content."""
    source_path: str
    source_name: str
    content_type: str  # text, table, image_description, mixed
    text_content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    extraction_timestamp: str = field(default_factory=lambda: datetime.now().isoformat())
    content_hash: str = ""
    
    def __post_init__(self):
        if not self.content_hash:
            self.content_hash = hashlib.md5(self.text_content.encode()).hexdigest()[:12]
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "source_path": self.source_path,
            "source_name": self.source_name,
            "content_type": self.content_type,
            "text_content": self.text_content,
            "metadata": self.metadata,
            "tables": self.tables,
            "images": self.images,
            "extraction_timestamp": self.extraction_timestamp,
            "content_hash": self.content_hash
        }


class ContentExtractor(ABC):
    """Base class for content extractors."""
    
    @property
    @abstractmethod
    def supported_extensions(self) -> List[str]:
        """List of file extensions this extractor handles."""
        pass
    
    @abstractmethod
    def extract(self, file_path: Path) -> ExtractedContent:
        """Extract content from file."""
        pass
    
    def can_handle(self, file_path: Path) -> bool:
        """Check if this extractor can handle the file."""
        return file_path.suffix.lower() in self.supported_extensions


class TextExtractor(ContentExtractor):
    """Extracts content from plain text files."""
    
    @property
    def supported_extensions(self) -> List[str]:
        return ['.txt', '.md', '.text', '.log', '.csv']
    
    def extract(self, file_path: Path) -> ExtractedContent:
        encodings = ['utf-8', 'utf-8-sig', 'cp1252', 'latin-1', 'iso-8859-7']
        
        content = None
        used_encoding = None
        
        for encoding in encodings:
            try:
                content = file_path.read_text(encoding=encoding)
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            raise ValueError(f"Could not decode {file_path} with any encoding")
        
        return ExtractedContent(
            source_path=str(file_path),
            source_name=file_path.name,
            content_type="text",
            text_content=content,
            metadata={
                "encoding": used_encoding,
                "size_bytes": file_path.stat().st_size,
                "line_count": content.count('\n') + 1
            }
        )


class PDFExtractor(ContentExtractor):
    """Extracts content from PDF files."""
    
    @property
    def supported_extensions(self) -> List[str]:
        return ['.pdf']
    
    def extract(self, file_path: Path) -> ExtractedContent:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            try:
                from pypdf import PdfReader
                return self._extract_with_pypdf(file_path)
            except ImportError:
                raise ImportError("Install PyMuPDF (fitz) or pypdf for PDF support")
        
        return self._extract_with_pymupdf(file_path)
    
    def _extract_with_pymupdf(self, file_path: Path) -> ExtractedContent:
        import fitz
        
        doc = fitz.open(str(file_path))
        text_parts = []
        images_info = []
        
        for page_num, page in enumerate(doc, 1):
            text_parts.append(f"--- Σελίδα {page_num} ---\n")
            text_parts.append(page.get_text())
            
            # Extract image metadata (not actual images for memory efficiency)
            for img_index, img in enumerate(page.get_images(full=True)):
                images_info.append({
                    "page": page_num,
                    "index": img_index,
                    "width": img[2],
                    "height": img[3]
                })
        
        doc.close()
        
        return ExtractedContent(
            source_path=str(file_path),
            source_name=file_path.name,
            content_type="mixed" if images_info else "text",
            text_content="\n".join(text_parts),
            images=images_info,
            metadata={
                "page_count": len(text_parts) // 2,
                "has_images": bool(images_info),
                "image_count": len(images_info)
            }
        )
    
    def _extract_with_pypdf(self, file_path: Path) -> ExtractedContent:
        from pypdf import PdfReader
        
        reader = PdfReader(str(file_path))
        text_parts = []
        
        for page_num, page in enumerate(reader.pages, 1):
            text_parts.append(f"--- Σελίδα {page_num} ---\n")
            text_parts.append(page.extract_text() or "")
        
        return ExtractedContent(
            source_path=str(file_path),
            source_name=file_path.name,
            content_type="text",
            text_content="\n".join(text_parts),
            metadata={"page_count": len(reader.pages)}
        )


class WordExtractor(ContentExtractor):
    """Extracts content from Word documents."""
    
    @property
    def supported_extensions(self) -> List[str]:
        return ['.docx', '.doc']
    
    def extract(self, file_path: Path) -> ExtractedContent:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Install python-docx for Word document support")
        
        if file_path.suffix.lower() == '.doc':
            logger.warning(f"Legacy .doc format may have limited support: {file_path}")
        
        doc = Document(str(file_path))
        
        text_parts = []
        tables_data = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                "index": table_idx,
                "rows": []
            }
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data["rows"].append(row_data)
            tables_data.append(table_data)
            
            # Also add table content as text
            text_parts.append(f"\n[Πίνακας {table_idx + 1}]")
            for row in table_data["rows"]:
                text_parts.append(" | ".join(row))
        
        return ExtractedContent(
            source_path=str(file_path),
            source_name=file_path.name,
            content_type="mixed" if tables_data else "text",
            text_content="\n".join(text_parts),
            tables=tables_data,
            metadata={
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(tables_data)
            }
        )


class ExcelExtractor(ContentExtractor):
    """Extracts content from Excel files."""
    
    @property
    def supported_extensions(self) -> List[str]:
        return ['.xlsx', '.xls', '.xlsm']
    
    def extract(self, file_path: Path) -> ExtractedContent:
        try:
            import openpyxl
        except ImportError:
            raise ImportError("Install openpyxl for Excel support")
        
        wb = openpyxl.load_workbook(str(file_path), data_only=True)
        
        text_parts = []
        tables_data = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"\n=== Φύλλο: {sheet_name} ===\n")
            
            table_data = {
                "sheet": sheet_name,
                "rows": []
            }
            
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    table_data["rows"].append(row_values)
                    text_parts.append(" | ".join(row_values))
            
            if table_data["rows"]:
                tables_data.append(table_data)
        
        wb.close()
        
        return ExtractedContent(
            source_path=str(file_path),
            source_name=file_path.name,
            content_type="table",
            text_content="\n".join(text_parts),
            tables=tables_data,
            metadata={
                "sheet_count": len(wb.sheetnames),
                "sheet_names": wb.sheetnames
            }
        )


class ImageExtractor(ContentExtractor):
    """Extracts metadata and optional OCR from images."""
    
    @property
    def supported_extensions(self) -> List[str]:
        return ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp']
    
    def extract(self, file_path: Path) -> ExtractedContent:
        metadata = {
            "file_size": file_path.stat().st_size
        }
        
        text_content = f"[Εικόνα: {file_path.name}]"
        
        # Try to get image dimensions
        try:
            from PIL import Image
            with Image.open(file_path) as img:
                metadata["width"] = img.width
                metadata["height"] = img.height
                metadata["format"] = img.format
                metadata["mode"] = img.mode
                text_content += f"\nΔιαστάσεις: {img.width}x{img.height}"
        except ImportError:
            logger.debug("PIL not available for image metadata")
        except Exception as e:
            logger.warning(f"Could not read image metadata: {e}")
        
        # Optional OCR support
        ocr_text = self._try_ocr(file_path)
        if ocr_text:
            text_content += f"\n\nΚείμενο από OCR:\n{ocr_text}"
            metadata["ocr_extracted"] = True
        
        return ExtractedContent(
            source_path=str(file_path),
            source_name=file_path.name,
            content_type="image_description",
            text_content=text_content,
            images=[{
                "path": str(file_path),
                "name": file_path.name,
                **metadata
            }],
            metadata=metadata
        )
    
    def _try_ocr(self, file_path: Path) -> Optional[str]:
        """Attempt OCR if pytesseract is available."""
        try:
            import pytesseract
            from PIL import Image
            
            with Image.open(file_path) as img:
                # Use Greek + English for OCR
                text = pytesseract.image_to_string(img, lang='ell+eng')
                return text.strip() if text.strip() else None
        except ImportError:
            return None
        except Exception as e:
            logger.debug(f"OCR failed: {e}")
            return None


class ContentExtractorRegistry:
    """Registry of content extractors with automatic format detection."""
    
    def __init__(self):
        self._extractors: List[ContentExtractor] = []
        self._register_defaults()
    
    def _register_defaults(self):
        """Register default extractors."""
        self._extractors = [
            TextExtractor(),
            PDFExtractor(),
            WordExtractor(),
            ExcelExtractor(),
            ImageExtractor(),
        ]
    
    def register(self, extractor: ContentExtractor):
        """Register a custom extractor."""
        self._extractors.insert(0, extractor)  # Custom extractors take priority
    
    def get_extractor(self, file_path: Path) -> Optional[ContentExtractor]:
        """Get appropriate extractor for file."""
        for extractor in self._extractors:
            if extractor.can_handle(file_path):
                return extractor
        return None
    
    def extract(self, file_path: Union[str, Path]) -> ExtractedContent:
        """Extract content from any supported file."""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        
        extractor = self.get_extractor(path)
        if not extractor:
            raise ValueError(f"No extractor available for {path.suffix}")
        
        logger.info(f"Extracting content from {path.name} using {type(extractor).__name__}")
        return extractor.extract(path)
    
    def extract_batch(self, file_paths: List[Union[str, Path]]) -> List[ExtractedContent]:
        """Extract content from multiple files."""
        results = []
        
        for file_path in file_paths:
            try:
                content = self.extract(file_path)
                results.append(content)
            except Exception as e:
                logger.error(f"Failed to extract {file_path}: {e}")
                # Create error placeholder
                results.append(ExtractedContent(
                    source_path=str(file_path),
                    source_name=Path(file_path).name,
                    content_type="error",
                    text_content=f"[Σφάλμα εξαγωγής: {str(e)}]",
                    metadata={"error": str(e)}
                ))
        
        return results
    
    @property
    def supported_extensions(self) -> List[str]:
        """Get all supported file extensions."""
        extensions = set()
        for extractor in self._extractors:
            extensions.update(extractor.supported_extensions)
        return sorted(extensions)


# Singleton registry
_registry: Optional[ContentExtractorRegistry] = None


def get_extractor_registry() -> ContentExtractorRegistry:
    """Get or create the extractor registry singleton."""
    global _registry
    if _registry is None:
        _registry = ContentExtractorRegistry()
    return _registry


def extract_content(file_path: Union[str, Path]) -> ExtractedContent:
    """Convenience function to extract content from a file."""
    return get_extractor_registry().extract(file_path)


def extract_batch(file_paths: List[Union[str, Path]]) -> List[ExtractedContent]:
    """Convenience function to extract content from multiple files."""
    return get_extractor_registry().extract_batch(file_paths)
# app/analysis/briefing_generator.py
"""
Intelligence briefing document generator.
Creates structured reports in Greek with proper citations and formatting.
"""

from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
from enum import Enum

from app.utils.logger import setup_logger
from app.config import BASE_DIR
from app.analysis.content_extractors import ExtractedContent
from app.analysis.intelligence_analyzer import AnalysisResult, IntelligencePattern

logger = setup_logger(__name__)


class ClassificationLevel(Enum):
    """Document classification levels."""
    UNCLASSIFIED = "ΑΔΙΑΒΑΘΜΗΤΟ"
    RESTRICTED = "ΠΕΡΙΟΡΙΣΜΕΝΗΣ ΧΡΗΣΗΣ"
    CONFIDENTIAL = "ΕΜΠΙΣΤΕΥΤΙΚΟ"
    SECRET = "ΑΠΟΡΡΗΤΟ"
    TOP_SECRET = "ΑΚΡΩΣ ΑΠΟΡΡΗΤΟ"


class ReportSection(Enum):
    """Standard briefing report sections."""
    EXECUTIVE_SUMMARY = "ΕΚΤΕΛΕΣΤΙΚΗ ΣΥΝΟΨΗ"
    SITUATION_OVERVIEW = "ΕΠΙΣΚΟΠΗΣΗ ΚΑΤΑΣΤΑΣΗΣ"
    KEY_FINDINGS = "ΚΥΡΙΑ ΕΥΡΗΜΑΤΑ"
    DETAILED_ANALYSIS = "ΑΝΑΛΥΤΙΚΗ ΑΝΑΛΥΣΗ"
    ENTITY_ANALYSIS = "ΑΝΑΛΥΣΗ ΟΝΤΟΤΗΤΩΝ"
    TIMELINE = "ΧΡΟΝΟΔΙΑΓΡΑΜΜΑ"
    CROSS_REFERENCES = "ΔΙΑΣΤΑΥΡΩΜΕΝΕΣ ΑΝΑΦΟΡΕΣ"
    SOURCES = "ΠΗΓΕΣ"
    APPENDIX = "ΠΑΡΑΡΤΗΜΑ"


@dataclass
class BriefingConfig:
    """Configuration for briefing generation."""
    classification: ClassificationLevel = ClassificationLevel.UNCLASSIFIED
    include_sources: bool = True
    include_timeline: bool = True
    include_entity_graph: bool = True
    include_confidence_scores: bool = True
    max_findings: int = 10
    language: str = "el"  # Greek
    date_format: str = "%d/%m/%Y %H:%M"
    output_format: str = "docx"  # docx, pdf, md


@dataclass
class Citation:
    """A citation reference."""
    id: int
    source_name: str
    source_path: str
    context: str
    page: Optional[int] = None
    
    def format_inline(self) -> str:
        """Format for inline citation."""
        return f"[{self.id}]"
    
    def format_reference(self) -> str:
        """Format for reference list."""
        ref = f"[{self.id}] {self.source_name}"
        if self.page:
            ref += f", σελ. {self.page}"
        return ref


class BriefingGenerator:
    """
    Generates comprehensive intelligence briefing documents.
    """
    
    def __init__(self, config: BriefingConfig = None, output_dir: Path = None):
        self.config = config or BriefingConfig()
        self.output_dir = output_dir or (BASE_DIR / "outputs" / "briefings")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._citations: List[Citation] = []
        self._citation_counter = 0
    
    def generate(
        self,
        analysis: AnalysisResult,
        documents: List[ExtractedContent],
        title: str = "Αναφορά Πληροφοριών"
    ) -> Path:
        """
        Generate complete briefing document.
        
        Returns path to generated document.
        """
        self._reset_citations()
        
        # Build document content
        content = self._build_content(analysis, documents, title)
        
        # Generate in requested format
        if self.config.output_format == "docx":
            return self._generate_docx(content, title)
        elif self.config.output_format == "pdf":
            return self._generate_pdf(content, title)
        elif self.config.output_format == "md":
            return self._generate_markdown(content, title)
        else:
            raise ValueError(f"Unsupported format: {self.config.output_format}")
    
    def _reset_citations(self):
        """Reset citation tracking."""
        self._citations = []
        self._citation_counter = 0
    
    def _add_citation(self, source_name: str, source_path: str, context: str, page: int = None) -> Citation:
        """Add a citation and return it."""
        self._citation_counter += 1
        citation = Citation(
            id=self._citation_counter,
            source_name=source_name,
            source_path=source_path,
            context=context[:100],
            page=page
        )
        self._citations.append(citation)
        return citation
    
    def _build_content(
        self,
        analysis: AnalysisResult,
        documents: List[ExtractedContent],
        title: str
    ) -> Dict[str, Any]:
        """Build structured content for the report."""
        content = {
            "title": title,
            "classification": self.config.classification.value,
            "generated_at": datetime.now().strftime(self.config.date_format),
            "sections": []
        }
        
        # Executive Summary
        content["sections"].append({
            "type": ReportSection.EXECUTIVE_SUMMARY,
            "title": ReportSection.EXECUTIVE_SUMMARY.value,
            "content": self._build_executive_summary(analysis, documents)
        })
        
        # Situation Overview
        content["sections"].append({
            "type": ReportSection.SITUATION_OVERVIEW,
            "title": ReportSection.SITUATION_OVERVIEW.value,
            "content": self._build_situation_overview(analysis, documents)
        })
        
        # Key Findings
        content["sections"].append({
            "type": ReportSection.KEY_FINDINGS,
            "title": ReportSection.KEY_FINDINGS.value,
            "content": self._build_key_findings(analysis)
        })
        
        # Detailed Analysis
        content["sections"].append({
            "type": ReportSection.DETAILED_ANALYSIS,
            "title": ReportSection.DETAILED_ANALYSIS.value,
            "content": self._build_detailed_analysis(analysis, documents)
        })
        
        # Entity Analysis
        if self.config.include_entity_graph and analysis.entity_graph:
            content["sections"].append({
                "type": ReportSection.ENTITY_ANALYSIS,
                "title": ReportSection.ENTITY_ANALYSIS.value,
                "content": self._build_entity_analysis(analysis)
            })
        
        # Timeline
        if self.config.include_timeline and analysis.timeline:
            content["sections"].append({
                "type": ReportSection.TIMELINE,
                "title": ReportSection.TIMELINE.value,
                "content": self._build_timeline(analysis)
            })
        
        # Cross-References
        if analysis.cross_references:
            content["sections"].append({
                "type": ReportSection.CROSS_REFERENCES,
                "title": ReportSection.CROSS_REFERENCES.value,
                "content": self._build_cross_references(analysis)
            })
        
        # Sources
        if self.config.include_sources:
            content["sections"].append({
                "type": ReportSection.SOURCES,
                "title": ReportSection.SOURCES.value,
                "content": self._build_sources(documents)
            })
        
        content["citations"] = self._citations
        
        return content
    
    def _build_executive_summary(
        self,
        analysis: AnalysisResult,
        documents: List[ExtractedContent]
    ) -> Dict[str, Any]:
        """Build executive summary section."""
        # Count high-priority patterns
        high_priority = [p for p in analysis.patterns if 'high_priority' in p.pattern_type]
        
        return {
            "summary_text": analysis.summary,
            "document_count": len(documents),
            "pattern_count": len(analysis.patterns),
            "high_priority_count": len(high_priority),
            "confidence_score": f"{analysis.confidence_score:.0%}" if self.config.include_confidence_scores else None,
            "analysis_date": analysis.analysis_timestamp
        }
    
    def _build_situation_overview(
        self,
        analysis: AnalysisResult,
        documents: List[ExtractedContent]
    ) -> Dict[str, Any]:
        """Build situation overview section."""
        # Categorize documents
        doc_categories = {}
        for doc in documents:
            cat = doc.content_type
            if cat not in doc_categories:
                doc_categories[cat] = []
            doc_categories[cat].append(doc.source_name)
        
        return {
            "document_breakdown": doc_categories,
            "total_documents": len(documents),
            "analysis_scope": "Πλήρης ανάλυση περιεχομένου και εξαγωγή μοτίβων"
        }
    
    def _build_key_findings(self, analysis: AnalysisResult) -> Dict[str, Any]:
        """Build key findings section."""
        return {
            "findings": analysis.key_findings[:self.config.max_findings],
            "priority_items": [
                {
                    "value": p.value,
                    "sources": p.sources,
                    "confidence": f"{p.confidence:.0%}"
                }
                for p in analysis.patterns
                if 'high_priority' in p.pattern_type
            ][:5]
        }
    
    def _build_detailed_analysis(
        self,
        analysis: AnalysisResult,
        documents: List[ExtractedContent]
    ) -> Dict[str, Any]:
        """Build detailed analysis section with citations."""
        # Group patterns by type
        patterns_by_type: Dict[str, List[IntelligencePattern]] = {}
        for pattern in analysis.patterns:
            ptype = pattern.pattern_type
            if ptype not in patterns_by_type:
                patterns_by_type[ptype] = []
            patterns_by_type[ptype].append(pattern)
        
        sections = []
        
        # Process each pattern type
        pattern_type_names = {
            "date": "Χρονικές Αναφορές",
            "location": "Γεωγραφικές Αναφορές",
            "entity": "Οντότητες & Πρόσωπα",
            "keyword_high_priority": "Υψηλής Προτεραιότητας",
            "keyword_entities": "Οργανωτικές Μονάδες",
            "keyword_actions": "Ενέργειες & Κινήσεις",
            "keyword_equipment": "Εξοπλισμός"
        }
        
        for ptype, patterns in patterns_by_type.items():
            type_name = pattern_type_names.get(ptype, ptype.replace("_", " ").title())
            
            items = []
            for p in patterns[:10]:  # Limit items per category
                # Add citation for each finding
                if p.sources:
                    doc = next((d for d in documents if d.source_name == p.sources[0]), None)
                    if doc:
                        citation = self._add_citation(
                            source_name=p.sources[0],
                            source_path=doc.source_path,
                            context=p.context_snippets[0] if p.context_snippets else ""
                        )
                        items.append({
                            "value": p.value,
                            "frequency": p.frequency,
                            "sources": p.sources,
                            "citation": citation.format_inline(),
                            "confidence": p.confidence
                        })
            
            if items:
                sections.append({
                    "category": type_name,
                    "items": items
                })
        
        return {"analysis_sections": sections}
    
    def _build_entity_analysis(self, analysis: AnalysisResult) -> Dict[str, Any]:
        """Build entity relationship analysis."""
        # Convert graph to list format
        relationships = []
        seen = set()
        
        for entity, related in analysis.entity_graph.items():
            for rel in related:
                pair = tuple(sorted([entity, rel]))
                if pair not in seen:
                    seen.add(pair)
                    relationships.append({
                        "entity1": entity,
                        "entity2": rel,
                        "relationship": "συσχετίζεται"
                    })
        
        return {
            "entity_count": len(analysis.entity_graph),
            "relationships": relationships[:20]  # Limit for readability
        }
    
    def _build_timeline(self, analysis: AnalysisResult) -> Dict[str, Any]:
        """Build timeline section."""
        return {
            "events": [
                {
                    "date": event["date"],
                    "sources": event["sources"],
                    "context": event["context"][:150] + "..." if len(event.get("context", "")) > 150 else event.get("context", "")
                }
                for event in analysis.timeline
            ]
        }
    
    def _build_cross_references(self, analysis: AnalysisResult) -> Dict[str, Any]:
        """Build cross-references section."""
        return {
            "references": [
                {
                    "source": cr.source_doc,
                    "target": cr.target_doc,
                    "type": cr.reference_type,
                    "common": cr.common_elements,
                    "score": f"{cr.relevance_score:.0%}"
                }
                for cr in analysis.cross_references
            ]
        }
    
    def _build_sources(self, documents: List[ExtractedContent]) -> Dict[str, Any]:
        """Build sources section."""
        return {
            "document_sources": [
                {
                    "name": doc.source_name,
                    "type": doc.content_type,
                    "path": doc.source_path,
                    "extracted_at": doc.extraction_timestamp
                }
                for doc in documents
            ],
            "citation_list": [c.format_reference() for c in self._citations]
        }
    
    def _generate_docx(self, content: Dict[str, Any], title: str) -> Path:
        """Generate Word document."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            raise ImportError("python-docx required for DOCX generation")
        
        doc = Document()
        
        # Classification header
        classification = doc.add_paragraph()
        classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = classification.add_run(content["classification"])
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        
        # Title
        title_para = doc.add_heading(content["title"], level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Ημερομηνία: {content['generated_at']}\n").italic = True
        
        # Add sections
        for section in content["sections"]:
            self._add_docx_section(doc, section)
        
        # Classification footer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(content["classification"])
        footer_run.bold = True
        footer_run.font.color.rgb = RGBColor(255, 0, 0)
        
        # Save
        filename = self._sanitize_filename(title) + ".docx"
        file_path = self.output_dir / filename
        doc.save(str(file_path))
        
        logger.info(f"Generated DOCX briefing: {file_path}")
        return file_path
    
    def _add_docx_section(self, doc, section: Dict[str, Any]):
        """Add a section to the Word document."""
        from docx.shared import Pt
        
        # Section heading
        doc.add_heading(section["title"], level=1)
        
        content = section["content"]
        section_type = section["type"]
        
        if section_type == ReportSection.EXECUTIVE_SUMMARY:
            doc.add_paragraph(content["summary_text"])
            
            summary_table = doc.add_table(rows=4, cols=2)
            summary_table.style = 'Table Grid'
            
            rows_data = [
                ("Έγγραφα που αναλύθηκαν", str(content["document_count"])),
                ("Μοτίβα που εντοπίστηκαν", str(content["pattern_count"])),
                ("Υψηλής προτεραιότητας", str(content["high_priority_count"])),
            ]
            if content.get("confidence_score"):
                rows_data.append(("Βαθμός εμπιστοσύνης", content["confidence_score"]))
            
            for i, (label, value) in enumerate(rows_data):
                if i < len(summary_table.rows):
                    summary_table.rows[i].cells[0].text = label
                    summary_table.rows[i].cells[1].text = value
        
        elif section_type == ReportSection.KEY_FINDINGS:
            for finding in content["findings"]:
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(finding)
            
            if content.get("priority_items"):
                doc.add_paragraph()
                doc.add_paragraph("Στοιχεία Υψηλής Προτεραιότητας:").bold = True
                for item in content["priority_items"]:
                    para = doc.add_paragraph(style='List Bullet')
                    para.add_run(f"{item['value']} - Πηγές: {', '.join(item['sources'])}")
        
        elif section_type == ReportSection.DETAILED_ANALYSIS:
            for analysis_section in content.get("analysis_sections", []):
                doc.add_heading(analysis_section["category"], level=2)
                
                for item in analysis_section["items"]:
                    para = doc.add_paragraph(style='List Bullet')
                    text = f"{item['value']}"
                    if item.get("frequency", 1) > 1:
                        text += f" (x{item['frequency']})"
                    text += f" {item.get('citation', '')}"
                    para.add_run(text)
        
        elif section_type == ReportSection.TIMELINE:
            for event in content.get("events", []):
                para = doc.add_paragraph()
                para.add_run(f"• {event['date']}: ").bold = True
                para.add_run(event.get("context", ""))
                if event.get("sources"):
                    para.add_run(f" [Πηγές: {', '.join(event['sources'])}]").italic = True
        
        elif section_type == ReportSection.SOURCES:
            doc.add_heading("Έγγραφα Πηγές", level=2)
            for src in content.get("document_sources", []):
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(f"{src['name']} ({src['type']})")
            
            if content.get("citation_list"):
                doc.add_heading("Αναφορές", level=2)
                for citation in content["citation_list"]:
                    doc.add_paragraph(citation)
        
        else:
            # Generic section handling
            if isinstance(content, dict):
                for key, value in content.items():
                    if isinstance(value, list):
                        for item in value:
                            doc.add_paragraph(str(item), style='List Bullet')
                    else:
                        doc.add_paragraph(f"{key}: {value}")
    
    def _generate_pdf(self, content: Dict[str, Any], title: str) -> Path:
        """Generate PDF document."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            raise ImportError("reportlab required for PDF generation")
        
        filename = self._sanitize_filename(title) + ".pdf"
        file_path = self.output_dir / filename
        
        doc = SimpleDocTemplate(str(file_path), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Classification
        class_style = ParagraphStyle(
            'Classification',
            parent=styles['Normal'],
            alignment=1,  # Center
            textColor=colors.red,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph(content["classification"], class_style))
        story.append(Spacer(1, 20))
        
        # Title
        story.append(Paragraph(content["title"], styles['Title']))
        story.append(Spacer(1, 12))
        
        # Date
        story.append(Paragraph(f"Ημερομηνία: {content['generated_at']}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Sections
        for section in content["sections"]:
            story.append(Paragraph(section["title"], styles['Heading1']))
            story.append(Spacer(1, 12))
            
            section_content = section["content"]
            if isinstance(section_content, dict):
                if "summary_text" in section_content:
                    story.append(Paragraph(section_content["summary_text"], styles['Normal']))
                elif "findings" in section_content:
                    for finding in section_content["findings"]:
                        story.append(Paragraph(f"• {finding}", styles['Normal']))
            
            story.append(Spacer(1, 12))
        
        # Footer classification
        story.append(Spacer(1, 30))
        story.append(Paragraph(content["classification"], class_style))
        
        doc.build(story)
        
        logger.info(f"Generated PDF briefing: {file_path}")
        return file_path
    
    def _generate_markdown(self, content: Dict[str, Any], title: str) -> Path:
        """Generate Markdown document."""
        lines = []
        
        # Classification
        lines.append(f"**{content['classification']}**")
        lines.append("")
        
        # Title
        lines.append(f"# {content['title']}")
        lines.append("")
        lines.append(f"*Ημερομηνία: {content['generated_at']}*")
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Sections
        for section in content["sections"]:
            lines.append(f"## {section['title']}")
            lines.append("")
            
            section_content = section["content"]
            if isinstance(section_content, dict):
                if "summary_text" in section_content:
                    lines.append(section_content["summary_text"])
                    lines.append("")
                
                if "findings" in section_content:
                    for finding in section_content["findings"]:
                        lines.append(f"- {finding}")
                    lines.append("")
                
                if "analysis_sections" in section_content:
                    for analysis in section_content["analysis_sections"]:
                        lines.append(f"### {analysis['category']}")
                        for item in analysis["items"]:
                            citation = item.get("citation", "")
                            lines.append(f"- {item['value']} {citation}")
                        lines.append("")
                
                if "document_sources" in section_content:
                    for src in section_content["document_sources"]:
                        lines.append(f"- **{src['name']}** ({src['type']})")
                    lines.append("")
                
                if "citation_list" in section_content:
                    lines.append("### Αναφορές")
                    for citation in section_content["citation_list"]:
                        lines.append(citation)
                    lines.append("")
            
            lines.append("")
        
        # Footer
        lines.append("---")
        lines.append(f"**{content['classification']}**")
        
        # Save
        filename = self._sanitize_filename(title) + ".md"
        file_path = self.output_dir / filename
        file_path.write_text("\n".join(lines), encoding='utf-8')
        
        logger.info(f"Generated Markdown briefing: {file_path}")
        return file_path
    
    def _sanitize_filename(self, name: str) -> str:
        """Sanitize filename for safe saving."""
        import re
        # Remove or replace invalid characters
        sanitized = re.sub(r'[<>:"/\\|?*]', '_', name)
        # Limit length
        sanitized = sanitized[:100]
        # Add timestamp for uniqueness
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{sanitized}_{timestamp}"


def generate_briefing(
    analysis: AnalysisResult,
    documents: List[ExtractedContent],
    title: str = "Αναφορά Πληροφοριών",
    config: BriefingConfig = None,
    output_format: str = "docx"
) -> Path:
    """Convenience function to generate a briefing document."""
    if config is None:
        config = BriefingConfig(output_format=output_format)
    else:
        config.output_format = output_format
    
    generator = BriefingGenerator(config)
    return generator.generate(analysis, documents, title)
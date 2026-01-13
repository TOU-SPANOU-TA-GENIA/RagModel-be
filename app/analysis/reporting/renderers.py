from abc import ABC, abstractmethod
from typing import Dict, Any
from pathlib import Path

class ReportRenderer(ABC):
    @abstractmethod
    def render(self, content: Dict[str, Any], output_path: Path) -> Path:
        pass

class DocxRenderer(ReportRenderer):
    def render(self, content: Dict[str, Any], output_path: Path) -> Path:
        # Import docx here to keep dependencies optional/clean
        from docx import Document
        doc = Document()
        doc.add_heading(content.get('title', 'Report'), 0)
        doc.add_paragraph(content.get('summary', ''))
        
        for section in content.get('sections', []):
            doc.add_heading(section['title'], 1)
            doc.add_paragraph(section['content'])
            
        doc.save(output_path)
        return output_path

# Add PDFRenderer, MarkdownRenderer here...